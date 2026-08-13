#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CONSTRUCTOR DE CV OPTIMIZADO PARA ATS

Toma el texto plano de un CV, lo reestructura con ayuda de Ollama y lo
reconstruye en un DOCX que cumple las reglas de formato que exigen los
sistemas ATS (Applicant Tracking Systems) y que además se lee bien cuando
un reclutador lo revisa a mano.
"""

import io
import json
import logging
import re
from typing import Any, Dict, List

logger = logging.getLogger(__name__)


# ==================== REGLAS DE FORMATO ATS ====================
# Estas constantes son la "hoja de estilo" del CV generado. Los ATS parsean
# mal las fuentes decorativas, las tablas, las columnas y todo lo que viva en
# encabezados o pies de página, así que el documento se arma sin nada de eso.

FUENTE = 'Calibri'          # Fuente sans-serif estándar, legible por cualquier parser
TAM_NOMBRE = 22             # pt - el nombre es lo único grande del documento
TAM_TITULAR = 11.5          # pt - el cargo objetivo, debajo del nombre
TAM_CONTACTO = 9.5          # pt
TAM_SECCION = 11            # pt - títulos de sección, en mayúscula y negrita
TAM_CUERPO = 10.5           # pt - texto normal; por debajo de 10 se vuelve ilegible impreso
TAM_META = 9.5              # pt - fechas y ubicación de cada puesto
MARGEN_CM = 1.9             # cm en los cuatro lados
INTERLINEADO = 1.12
ANCHO_UTIL_CM = 21.0 - (2 * MARGEN_CM)   # A4 menos los márgenes: dónde cae el tabulador derecho

COLOR_TEXTO = (0x1a, 0x1a, 0x1a)
COLOR_SECUNDARIO = (0x5a, 0x5a, 0x5a)    # fechas, ubicación, datos de contacto
COLOR_LINEA = '9a9a9a'

SECCIONES_ORDEN = [
    ('resumen', 'PERFIL PROFESIONAL'),
    ('experiencia', 'EXPERIENCIA PROFESIONAL'),
    ('educacion', 'EDUCACIÓN'),
    ('habilidades_tecnicas', 'HABILIDADES TÉCNICAS'),
    ('habilidades_blandas', 'HABILIDADES BLANDAS'),
    ('certificaciones', 'CERTIFICACIONES'),
    ('idiomas', 'IDIOMAS'),
]

# Estructura vacía: todo lo que el resto del módulo espera encontrar.
ESQUEMA_VACIO = {
    'nombre': '',
    'titular': '',
    'contacto': {'email': '', 'telefono': '', 'ciudad': '', 'linkedin': ''},
    'resumen': '',
    'experiencia': [],
    'educacion': [],
    'habilidades_tecnicas': [],
    'habilidades_blandas': [],
    'certificaciones': [],
    'idiomas': [],
    'referencias': [],
}


PROMPT_REESTRUCTURA = """Eres un experto en reclutamiento y en optimización de CVs para sistemas ATS.

Recibes el texto crudo de un CV. Tu tarea es reescribirlo y devolverlo estructurado en JSON.

REGLA MÁS IMPORTANTE: no puedes inventar NADA.
Este CV lo va a enviar una persona real a una empresa real. Un dato inventado
es una mentira en su hoja de vida. Ante la duda, deja el campo vacío.

Prohibido específicamente:
- Inventar años de experiencia ("más de 5 años") si el CV no los dice.
- Inventar logros o funciones que el CV no menciona. Si un puesto no describe
  funciones, deja "logros" como lista vacía []. Es correcto dejarlo vacío.
- Inventar habilidades, herramientas o certificaciones que no aparezcan.
- Convertir a alguien en líder ("Lideré", "Coordiné equipos") si el CV no lo dice.

REGLAS OBLIGATORIAS:
1. NO inventes datos. Si un dato no está en el CV original, deja el campo vacío ("" o []).
   Nunca inventes cifras, empresas, fechas, títulos ni certificaciones.
2. Reescribe cada logro empezando con un VERBO DE ACCIÓN en pasado
   (Lideré, Desarrollé, Implementé, Reduje, Automaticé, Coordiné...).
3. Cuantifica SOLO con números que ya aparezcan en el texto original.
4. Cada logro es una línea de máximo 2 renglones, sin punto final innecesario.
5. Escribe los nombres de tecnologías tal como los busca un ATS
   (ejemplo: "Python", "SQL", "Power BI", "Excel avanzado").
6. El resumen tiene entre 2 y 4 líneas y menciona años de experiencia y área.
7. Todo en español, salvo nombres propios de tecnologías.
8. En "nombre" va la persona del CV, nunca el título del documento
   ("HOJA DE VIDA", "CURRICULUM VITAE") ni el nombre de una empresa.
9. NO incluyas datos de contacto de terceros de la empresa (jefe directo,
   recursos humanos, teléfono de la empresa, NIT): no van en un CV.
   Las referencias personales sí, en su propio campo.
{contexto_cargo}
Devuelve ÚNICAMENTE este JSON, sin texto adicional:
{{
  "nombre": "",
  "titular": "cargo o perfil profesional en una línea",
  "contacto": {{"email": "", "telefono": "", "ciudad": "", "linkedin": ""}},
  "resumen": "",
  "experiencia": [
    {{"cargo": "", "empresa": "", "ubicacion": "", "inicio": "", "fin": "", "logros": ["", ""]}}
  ],
  "educacion": [{{"titulo": "", "institucion": "", "anio": ""}}],
  "habilidades_tecnicas": [],
  "habilidades_blandas": [],
  "certificaciones": [],
  "idiomas": [],
  "referencias": [{{"nombre": "", "cargo": "", "telefono": ""}}]
}}

CV ORIGINAL:
{cv_text}
"""


def _texto_de_item(valor: Any) -> str:
    """Convierte a texto un elemento de una lista simple.

    Algunos modelos devuelven objetos donde el esquema pide strings
    (por ejemplo {"nombre": "Inglés", "nivel": "Básico"} en idiomas).
    Sin esto, el CV terminaría mostrando el diccionario en crudo.
    """
    if isinstance(valor, dict):
        partes = [str(v).strip() for v in valor.values() if str(v or '').strip()]
        return ' - '.join(partes)
    if isinstance(valor, (list, tuple)):
        return ' '.join(_texto_de_item(v) for v in valor).strip()
    return str(valor or '').strip()


def _normalizar(data: Any) -> Dict[str, Any]:
    """Rellena los campos faltantes y fuerza los tipos que espera el renderizador.

    El modelo a veces devuelve una lista donde se espera un string, o se salta
    campos completos. Sin esto, el renderizado revienta en runtime.
    """
    resultado = json.loads(json.dumps(ESQUEMA_VACIO))  # copia profunda

    if not isinstance(data, dict):
        return resultado

    for campo in ('nombre', 'titular', 'resumen'):
        resultado[campo] = _texto_de_item(data.get(campo, ''))

    contacto = data.get('contacto') or {}
    if isinstance(contacto, dict):
        for campo in resultado['contacto']:
            resultado['contacto'][campo] = str(contacto.get(campo, '') or '').strip()

    for campo in ('habilidades_tecnicas', 'habilidades_blandas', 'certificaciones', 'idiomas'):
        valor = data.get(campo) or []
        if isinstance(valor, str):
            valor = [v.strip() for v in re.split(r'[,;\n]', valor)]
        if isinstance(valor, list):
            resultado[campo] = [t for t in (_texto_de_item(v) for v in valor) if t]

    experiencia = data.get('experiencia') or []
    if isinstance(experiencia, list):
        for item in experiencia:
            if not isinstance(item, dict):
                continue
            logros = item.get('logros') or []
            if isinstance(logros, str):
                logros = [logros]
            resultado['experiencia'].append({
                'cargo': str(item.get('cargo', '') or '').strip(),
                'empresa': str(item.get('empresa', '') or '').strip(),
                'ubicacion': str(item.get('ubicacion', '') or '').strip(),
                'inicio': str(item.get('inicio', '') or '').strip(),
                'fin': str(item.get('fin', '') or '').strip(),
                'logros': [t for t in (_texto_de_item(l) for l in logros) if t],
            })

    educacion = data.get('educacion') or []
    if isinstance(educacion, list):
        for item in educacion:
            if not isinstance(item, dict):
                continue
            resultado['educacion'].append({
                'titulo': str(item.get('titulo', '') or '').strip(),
                'institucion': str(item.get('institucion', '') or '').strip(),
                'anio': str(item.get('anio', '') or '').strip(),
            })

    referencias = data.get('referencias') or []
    if isinstance(referencias, list):
        for item in referencias:
            if isinstance(item, dict):
                entrada = {
                    'nombre': str(item.get('nombre', '') or '').strip(),
                    'cargo': str(item.get('cargo', '') or '').strip(),
                    'telefono': str(item.get('telefono', '') or '').strip(),
                }
            else:
                entrada = {'nombre': _texto_de_item(item), 'cargo': '', 'telefono': ''}
            if any(entrada.values()):
                resultado['referencias'].append(entrada)

    return resultado


def estructurar_con_ollama(bot, cv_text: str, cargo_objetivo: str = '') -> Dict[str, Any]:
    """Pide a Ollama que reestructure el CV y devuelve el JSON normalizado.

    Lanza excepción si Ollama falla; el que llama decide el fallback.
    """
    contexto_cargo = ''
    if cargo_objetivo:
        contexto_cargo = (
            f'8. El CV apunta al cargo de "{cargo_objetivo}". Prioriza y redacta '
            'las habilidades y logros que sean relevantes para ese cargo, pero sin '
            'inventar experiencia que no exista.\n'
        )

    prompt = PROMPT_REESTRUCTURA.format(
        contexto_cargo=contexto_cargo,
        cv_text=cv_text[:12000],  # límite defensivo: CVs muy largos desbordan el contexto
    )

    # Temperatura muy baja: con este prompt, cuanto menos "creativo" sea el
    # modelo, menos datos se inventa.
    respuesta = bot.call_ollama(prompt, fmt='json', options={'num_ctx': 8192, 'temperature': 0.1})

    try:
        return _normalizar(json.loads(respuesta))
    except json.JSONDecodeError:
        # Algunos modelos envuelven el JSON en texto o en un bloque markdown
        match = re.search(r'\{.*\}', respuesta, re.DOTALL)
        if match:
            return _normalizar(json.loads(match.group()))
        raise ValueError('Ollama no devolvió un JSON válido')


# ==================== FALLBACK SIN IA ====================

_RE_EMAIL = re.compile(r'[\w.+-]+@[\w-]+\.[\w.]+')
_RE_TELEFONO = re.compile(r'(?:\+\d{1,3}[\s-]?)?(?:\(\d{2,3}\)[\s-]?)?\d{3}[\s-]?\d{3}[\s-]?\d{2,4}')
_RE_LINKEDIN = re.compile(r'(?:linkedin\.com/in/|linkedin:\s*)([\w-]+)', re.IGNORECASE)

_ENCABEZADOS = {
    'resumen': ('perfil', 'resumen', 'objetivo', 'acerca de', 'sobre mi', 'sobre mí', 'presentacion', 'presentación'),
    'experiencia': ('experiencia', 'trayectoria', 'historial laboral', 'experiencia laboral'),
    'educacion': ('educacion', 'educación', 'formacion', 'formación', 'estudios', 'academic'),
    'habilidades_tecnicas': ('habilidades', 'competencias', 'skills', 'conocimientos', 'tecnologias', 'tecnologías', 'aptitudes'),
    'certificaciones': ('certificacion', 'certificación', 'certificaciones', 'cursos', 'diplomados'),
    'idiomas': ('idioma', 'idiomas', 'languages'),
    'referencias': ('referencia', 'referencias'),
    'datos_personales': ('datos personales', 'informacion personal', 'información personal', 'datos generales'),
}

# Títulos genéricos del documento: nunca son el nombre de la persona
_TITULOS_DOCUMENTO = (
    'hoja de vida', 'curriculum vitae', 'currículum vitae', 'curriculum',
    'currículum', 'cv', 'resume', 'perfil profesional', 'datos personales',
)

# Etiquetas de los CV con formato "Clave: valor", habituales en Colombia
_ETIQUETAS_CARGO = ('cargo', 'puesto', 'ocupacion', 'ocupación')
_ETIQUETAS_FECHA = ('tiempo servido', 'periodo', 'período', 'fecha', 'fechas', 'duracion', 'duración')
_ETIQUETAS_EMPRESA = ('empresa', 'compania', 'compañia', 'compañía', 'organizacion', 'organización')
_ETIQUETAS_TELEFONO = ('cel', 'celular', 'telefono', 'teléfono', 'tel', 'movil', 'móvil', 'contacto')
_ETIQUETAS_INSTITUCION = ('colegio', 'universidad', 'institucion', 'institución', 'centro')

# Datos de contacto de terceros y referencias internas de la empresa. No van en
# un CV: no aportan al reclutador y exponen datos de otras personas.
_ETIQUETAS_DESCARTAR = (
    'jefe directo', 'jefe inmediato', 'jefe', 'celular jefe', 'recursos humanos',
    'celular empresa', 'telefono empresa', 'teléfono empresa', 'nit', 'direccion empresa',
)


_TODOS_LOS_ALIAS = tuple(alias for grupo in _ENCABEZADOS.values() for alias in grupo)

# Palabras que introducen la fecha de finalización de un estudio
_PREFIJOS_FIN_ESTUDIO = ('finalizado', 'finalizada', 'graduado', 'graduada',
                         'culminado', 'culminada', 'terminado', 'terminada', 'año')


def _es_titulo_documento(linea: str) -> bool:
    limpia = re.sub(r'[^\w\s]', '', linea.lower()).strip()
    return limpia in _TITULOS_DOCUMENTO


def _separar_etiqueta(linea: str):
    """Divide 'Cargo: Operaria' en ('cargo', 'Operaria'). Devuelve None si no aplica."""
    if ':' not in linea:
        return None
    etiqueta, _, valor = linea.partition(':')
    etiqueta = etiqueta.strip().lower()
    # Una etiqueta real es corta; si no, seguramente es una frase con dos puntos
    if not etiqueta or len(etiqueta) > 30 or len(etiqueta.split()) > 4:
        return None
    return re.sub(r'[^\w\s]', '', etiqueta).strip(), valor.strip()


def _es_encabezado_seccion(linea: str) -> bool:
    """True si la línea es el título de una sección ('EXPERIENCIA PROFESIONAL')."""
    limpia = re.sub(r'[^\w\sáéíóúñ]', '', linea.lower()).strip()
    if len(limpia) > 40:
        return False
    return any(limpia.startswith(alias) for alias in _TODOS_LOS_ALIAS)


def _parece_nombre_persona(linea: str) -> bool:
    """Una línea con pinta de nombre propio: 2 a 5 palabras, sin datos ni signos."""
    if not linea or _es_titulo_documento(linea) or _es_encabezado_seccion(linea):
        return False
    if any(c.isdigit() for c in linea) or '@' in linea or ':' in linea:
        return False
    palabras = linea.split()
    if not 2 <= len(palabras) <= 5:
        return False
    # Todas las palabras empiezan en mayúscula, o la línea entera va en mayúsculas
    return linea.isupper() or all(p[:1].isupper() for p in palabras if len(p) > 2)


def _detectar_nombre(candidatas: List[str], email: str = '') -> str:
    """Elige, entre las líneas dadas, la que mejor parece el nombre del titular.

    Cuando hay varias candidatas se prefiere la que comparte palabras con el
    email, que suele construirse a partir del nombre.
    """
    nombres = [l for l in candidatas if _parece_nombre_persona(l)]
    if not nombres:
        return ''

    usuario = email.split('@')[0].lower() if email else ''
    # Se quitan tildes para poder comparar 'gutiérrez' con 'gutierrez'
    tabla = str.maketrans('áéíóúñ', 'aeioun')
    usuario = usuario.translate(tabla)

    def puntuar(linea):
        if not usuario:
            return 0
        palabras = [p.translate(tabla) for p in linea.lower().split() if len(p) > 2]
        return sum(1 for p in palabras if p in usuario)

    mejor = max(nombres, key=puntuar)
    return mejor.title() if mejor.isupper() else mejor


def estructurar_sin_ia(cv_text: str) -> Dict[str, Any]:
    """Reestructura el CV troceándolo por encabezados, cuando Ollama no está disponible.

    No reescribe el contenido —eso solo lo hace el LLM— pero sí produce un
    documento con el formato ATS correcto, que ya es la mitad del problema.
    """
    datos = json.loads(json.dumps(ESQUEMA_VACIO))
    lineas = [l.strip() for l in cv_text.splitlines()]

    email = _RE_EMAIL.search(cv_text)
    if email:
        datos['contacto']['email'] = email.group()

    linkedin = _RE_LINKEDIN.search(cv_text)
    if linkedin:
        datos['contacto']['linkedin'] = f'linkedin.com/in/{linkedin.group(1)}'

    # El teléfono se busca solo en la cabecera para no capturar años ni cifras de logros
    telefono = _RE_TELEFONO.search('\n'.join(lineas[:15]))
    if telefono and len(re.sub(r'\D', '', telefono.group())) >= 7:
        datos['contacto']['telefono'] = telefono.group().strip()

    # Trocear el cuerpo por encabezados de sección
    seccion_actual = None
    buffer: Dict[str, List[str]] = {clave: [] for clave in _ENCABEZADOS}
    cabecera: List[str] = []   # lo que aparece antes de la primera sección

    for linea in lineas:
        if not linea:
            continue
        limpia = re.sub(r'[^\w\sáéíóúñ]', '', linea.lower()).strip()
        detectada = None
        for clave, alias in _ENCABEZADOS.items():
            # Un encabezado es corto y empieza con alguna de sus palabras clave
            if len(limpia) <= 40 and any(limpia.startswith(a) for a in alias):
                detectada = clave
                break
        if detectada:
            seccion_actual = detectada
            continue
        if seccion_actual:
            buffer[seccion_actual].append(linea)
        else:
            cabecera.append(linea)

    # El nombre se busca en la cabecera y, si allí solo hay un título genérico
    # como "HOJA DE VIDA", en los datos personales y las referencias, donde
    # muchos CV repiten el nombre completo junto a la cédula. Nunca se busca en
    # experiencia: ahí las líneas en mayúsculas son nombres de empresas.
    datos['nombre'] = _detectar_nombre(
        cabecera + buffer['datos_personales'] + buffer['referencias'],
        datos['contacto']['email'],
    )

    datos['resumen'] = ' '.join(buffer['resumen'])[:600]

    for clave in ('habilidades_tecnicas', 'certificaciones', 'idiomas'):
        items: List[str] = []
        for linea in buffer[clave]:
            items.extend(re.split(r'[,;•·|]|\s{3,}', linea))
        datos[clave] = [i.strip(' -•\t') for i in items if len(i.strip(' -•\t')) > 1][:20]

    datos['experiencia'] = _trocear_experiencia(buffer['experiencia'])
    datos['educacion'] = _trocear_educacion(buffer['educacion'])
    datos['referencias'] = _trocear_referencias(buffer['referencias'])

    # Sin titular, la cabecera queda con el nombre suelto y el CV se ve
    # incompleto. Se usa el cargo más reciente, que ya está en el CV.
    if datos['experiencia']:
        datos['titular'] = datos['experiencia'][0].get('cargo', '')

    # En muchos CV el nombre completo aparece dentro del bloque de referencias,
    # junto a la cédula. Si el titular se coló ahí, se saca de la lista y sus
    # datos de contacto se aprovechan para la cabecera.
    if datos['nombre']:
        propio = datos['nombre'].lower()
        restantes = []
        for ref in datos['referencias']:
            if ref['nombre'].lower() == propio:
                if ref['telefono'] and not datos['contacto']['telefono']:
                    datos['contacto']['telefono'] = ref['telefono']
            else:
                restantes.append(ref)
        datos['referencias'] = restantes

    return datos


def _cerrar_bloque(bloque, destino):
    """Guarda el puesto en curso solo si tiene sustancia.

    Una empresa suelta, sin cargo, sin funciones y sin fechas, casi siempre es
    una línea residual (la ciudad al final de un bloque, por ejemplo), no un
    puesto de trabajo.
    """
    if not bloque:
        return
    tiene_contenido = bool(bloque['cargo'] or bloque['logros'])
    empresa_con_fechas = bool(bloque['empresa'] and (bloque['inicio'] or bloque['fin']))
    if tiene_contenido or empresa_con_fechas:
        destino.append(bloque)


def _nuevo_bloque():
    return {'cargo': '', 'empresa': '', 'ubicacion': '', 'inicio': '', 'fin': '', 'logros': []}


def _fechas_de(texto: str):
    """Devuelve (inicio, fin) a partir de los años que aparezcan en el texto."""
    anios = re.findall(r'(?:19|20)\d{2}', texto)
    if not anios:
        return '', ''
    if len(anios) > 1:
        return anios[0], anios[1]
    # Un solo año con guion suelto suele significar "hasta hoy"
    if re.search(r'(actual|presente|hoy)', texto, re.IGNORECASE) or texto.rstrip().endswith('-'):
        return anios[0], 'Actual'
    return anios[0], ''


def _trocear_experiencia(lineas: List[str]) -> List[Dict[str, Any]]:
    """Agrupa la sección de experiencia en puestos.

    Contempla el formato habitual en los CV colombianos, donde cada puesto se
    describe con pares "Cargo: ...", "Tiempo servido: ...". Antes cada una de
    esas líneas se tomaba como un puesto nuevo y el documento salía como una
    lista suelta de frases en negrita.
    """
    puestos: List[Dict[str, Any]] = []
    bloque = _nuevo_bloque()

    for linea in lineas:
        if len(linea) < 2:
            continue

        # Viñeta explícita: es una función o un logro del puesto en curso
        if re.match(r'^[-•*·▪]', linea):
            bloque['logros'].append(linea.lstrip(' -•*·▪\t'))
            continue

        etiquetado = _separar_etiqueta(linea)
        if etiquetado:
            etiqueta, valor = etiquetado

            if etiqueta in _ETIQUETAS_DESCARTAR:
                continue

            if etiqueta in _ETIQUETAS_CARGO:
                # Un segundo "Cargo:" significa que empezó otro puesto
                if bloque['cargo']:
                    _cerrar_bloque(bloque, puestos)
                    bloque = _nuevo_bloque()
                bloque['cargo'] = valor
                continue

            if etiqueta in _ETIQUETAS_FECHA:
                inicio, fin = _fechas_de(valor)
                bloque['inicio'] = bloque['inicio'] or inicio
                bloque['fin'] = bloque['fin'] or fin
                continue

            if etiqueta in _ETIQUETAS_EMPRESA:
                bloque['empresa'] = valor
                continue

            if etiqueta in _ETIQUETAS_TELEFONO:
                continue  # teléfonos de la empresa: no van en el CV

            # Cualquier otra etiqueta con contenido pasa a ser una función
            if valor:
                bloque['logros'].append(f'{etiqueta.capitalize()}: {valor}')
            continue

        # Línea suelta con años: completa las fechas del puesto en curso
        inicio, fin = _fechas_de(linea)
        if inicio and len(re.sub(r'[\d\s\-–—/]', '', linea)) <= 8:
            bloque['inicio'] = bloque['inicio'] or inicio
            bloque['fin'] = bloque['fin'] or fin
            continue

        # Línea sin etiqueta: abre un puesto nuevo
        if bloque['empresa'] or bloque['cargo'] or bloque['logros']:
            _cerrar_bloque(bloque, puestos)
            bloque = _nuevo_bloque()

        # Se quita el rango de fechas para que no acabe dentro del cargo o de
        # la empresa, y se separa el formato "Cargo - Empresa" si viene junto
        sin_fechas = re.sub(
            r'[|(\[]?\s*(?:19|20)\d{2}\s*[-–—/]?\s*'
            r'(?:(?:19|20)\d{2}|actual|presente|present)?\s*[)\]]?\s*$',
            '', linea, flags=re.IGNORECASE).strip(' -–—|,')

        partes = re.split(r'\s+[|–—-]\s+', sin_fechas, maxsplit=1)
        if len(partes) == 2:
            bloque['cargo'], bloque['empresa'] = partes[0].strip(), partes[1].strip()
        else:
            bloque['empresa'] = sin_fechas

        if inicio:
            bloque['inicio'], bloque['fin'] = inicio, fin

    _cerrar_bloque(bloque, puestos)
    return puestos


def _trocear_educacion(lineas: List[str]) -> List[Dict[str, Any]]:
    """Agrupa la sección de educación en títulos con institución y año."""
    estudios: List[Dict[str, Any]] = []
    actual = None

    for linea in lineas:
        if len(linea) < 3:
            continue

        etiquetado = _separar_etiqueta(linea)
        if etiquetado and actual:
            etiqueta, valor = etiquetado
            if etiqueta in _ETIQUETAS_INSTITUCION:
                actual['institucion'] = valor
                continue

        anio = re.search(r'(?:19|20)\d{2}', linea)
        primera_palabra = linea.split()[0].lower().strip('.,:')

        # "Finalizado 06 de diciembre de 1999" completa el estudio anterior,
        # igual que una línea que es casi solo una fecha
        if actual and anio and (primera_palabra in _PREFIJOS_FIN_ESTUDIO
                                or len(re.sub(r'[^A-Za-zÁÉÍÓÚáéíóúÑñ]', '', linea)) <= 12):
            actual['anio'] = actual['anio'] or anio.group()
            continue

        # Una sola palabra sin fecha suele ser la ciudad: no es un título
        if actual and not anio and len(linea.split()) == 1:
            continue

        if actual:
            estudios.append(actual)
        actual = {
            'titulo': re.sub(r'\s*\(?(?:19|20)\d{2}\)?\s*$', '', linea.strip(' -•\t')).strip(),
            'institucion': '',
            'anio': anio.group() if anio else '',
        }

    if actual:
        estudios.append(actual)
    return estudios


def _trocear_referencias(lineas: List[str]) -> List[Dict[str, Any]]:
    """Agrupa las referencias en nombre, cargo y teléfono."""
    referencias: List[Dict[str, Any]] = []
    actual = None

    for linea in lineas:
        if len(linea) < 3:
            continue

        etiquetado = _separar_etiqueta(linea)
        if etiquetado:
            etiqueta, valor = etiquetado
            if etiqueta in _ETIQUETAS_TELEFONO and actual:
                actual['telefono'] = valor
            continue

        if _parece_nombre_persona(linea):
            if actual:
                referencias.append(actual)
            actual = {'nombre': linea.title() if linea.isupper() else linea,
                      'cargo': '', 'telefono': ''}
        elif actual and not actual['cargo']:
            actual['cargo'] = linea

    if actual:
        referencias.append(actual)
    return referencias


# ==================== RENDERIZADO DOCX ====================

def construir_docx(datos: Dict[str, Any]) -> bytes:
    """Arma el DOCX aplicando las reglas de formato ATS. Devuelve los bytes.

    El diseño busca que un reclutador distinga de un vistazo el nombre, cada
    puesto y sus fechas, sin recurrir a nada que un ATS no sepa leer: no hay
    tablas, ni columnas, ni imágenes. La jerarquía se consigue solo con
    tamaños, grosores, espaciado y tabuladores.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()

    # Fuente base de todo el documento
    normal = doc.styles['Normal']
    normal.font.name = FUENTE
    normal.font.size = Pt(TAM_CUERPO)
    normal.font.color.rgb = RGBColor(*COLOR_TEXTO)
    normal.paragraph_format.line_spacing = INTERLINEADO
    normal.paragraph_format.space_after = Pt(3)

    for seccion in doc.sections:
        seccion.top_margin = Cm(MARGEN_CM)
        seccion.bottom_margin = Cm(MARGEN_CM)
        seccion.left_margin = Cm(MARGEN_CM)
        seccion.right_margin = Cm(MARGEN_CM)

    def _formatear(run, tam, negrita=False, cursiva=False, color=COLOR_TEXTO, espaciado_letras=None):
        run.font.name = FUENTE
        run.font.size = Pt(tam)
        run.font.bold = negrita
        run.font.italic = cursiva
        run.font.color.rgb = RGBColor(*color)
        if espaciado_letras:
            # El espaciado entre letras se mide en veinteavos de punto y da al
            # nombre y a los títulos un aire más cuidado sin cambiar la fuente
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:val'), str(int(espaciado_letras * 20)))
            run._element.get_or_add_rPr().append(spacing)
        return run

    def parrafo(texto='', tam=TAM_CUERPO, negrita=False, cursiva=False, alineacion=None,
                espacio_antes=0, espacio_despues=3, color=COLOR_TEXTO, espaciado_letras=None,
                sangria=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(espacio_antes)
        p.paragraph_format.space_after = Pt(espacio_despues)
        p.paragraph_format.line_spacing = INTERLINEADO
        if alineacion is not None:
            p.alignment = alineacion
        if sangria is not None:
            p.paragraph_format.left_indent = Cm(sangria)
        if texto:
            _formatear(p.add_run(texto), tam, negrita, cursiva, color, espaciado_letras)
        return p

    def linea_horizontal(p, grosor='6', color=COLOR_LINEA, espacio='2'):
        """Borde inferior de párrafo. El ATS lo ignora; el ojo lo agradece."""
        borde = OxmlElement('w:pBdr')
        linea = OxmlElement('w:bottom')
        linea.set(qn('w:val'), 'single')
        linea.set(qn('w:sz'), grosor)
        linea.set(qn('w:space'), espacio)
        linea.set(qn('w:color'), color)
        borde.append(linea)
        p._p.get_or_add_pPr().append(borde)
        return p

    def titulo_seccion(texto):
        p = parrafo(texto.upper(), tam=TAM_SECCION, negrita=True,
                    espacio_antes=11, espacio_despues=5, espaciado_letras=0.8)
        return linea_horizontal(p, grosor='4', espacio='3')

    def puesto(titulo, derecha=''):
        """Cargo a la izquierda y fechas alineadas a la derecha.

        Se usa un tabulador derecho al final del ancho útil. Es texto plano,
        así que el ATS lo lee en el mismo orden en que se ve.
        """
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = INTERLINEADO
        p.paragraph_format.tab_stops.add_tab_stop(Cm(ANCHO_UTIL_CM), WD_TAB_ALIGNMENT.RIGHT)
        _formatear(p.add_run(titulo), TAM_CUERPO, negrita=True)
        if derecha:
            _formatear(p.add_run('\t' + derecha), TAM_META, color=COLOR_SECUNDARIO)
        return p

    def vineta(texto):
        # 'List Bullet' es una lista nativa de Word: los parsers la leen como
        # viñeta real, a diferencia de un guion escrito a mano.
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = INTERLINEADO
        p.paragraph_format.left_indent = Cm(0.6)
        _formatear(p.add_run(texto), TAM_CUERPO)
        return p

    # --- Encabezado ---
    if datos.get('nombre'):
        parrafo(datos['nombre'].upper(), tam=TAM_NOMBRE, negrita=True,
                alineacion=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=1,
                espaciado_letras=1.2)

    if datos.get('titular'):
        parrafo(datos['titular'], tam=TAM_TITULAR, alineacion=WD_ALIGN_PARAGRAPH.CENTER,
                espacio_despues=3, color=COLOR_SECUNDARIO)

    contacto = datos.get('contacto', {})
    partes_contacto = [contacto.get(c) for c in ('email', 'telefono', 'ciudad', 'linkedin')]
    partes_contacto = [p for p in partes_contacto if p]
    if partes_contacto:
        p = parrafo('  •  '.join(partes_contacto), tam=TAM_CONTACTO,
                    alineacion=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=2,
                    color=COLOR_SECUNDARIO)
        linea_horizontal(p, grosor='8', color='444444', espacio='6')

    # --- Cuerpo ---
    if datos.get('resumen'):
        titulo_seccion('Perfil profesional')
        parrafo(datos['resumen'], alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY, espacio_despues=2)

    if datos.get('experiencia'):
        titulo_seccion('Experiencia profesional')
        for exp in datos['experiencia']:
            puesto(exp.get('cargo') or exp.get('empresa', ''), _rango_fechas(exp))

            # La empresa va en su propia línea, en cursiva: separa el "qué
            # hacía" del "dónde" sin competir con el cargo.
            subtitulo = [x for x in (exp.get('empresa') if exp.get('cargo') else '',
                                     exp.get('ubicacion')) if x]
            if subtitulo:
                parrafo(' — '.join(subtitulo), tam=TAM_META, cursiva=True,
                        color=COLOR_SECUNDARIO, espacio_despues=3)

            for logro in exp.get('logros', []):
                vineta(logro)

    if datos.get('educacion'):
        titulo_seccion('Educación')
        for edu in datos['educacion']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = INTERLINEADO
            p.paragraph_format.tab_stops.add_tab_stop(Cm(ANCHO_UTIL_CM), WD_TAB_ALIGNMENT.RIGHT)
            _formatear(p.add_run(edu.get('titulo', '')), TAM_CUERPO, negrita=True)
            if edu.get('anio'):
                _formatear(p.add_run('\t' + edu['anio']), TAM_META, color=COLOR_SECUNDARIO)
            if edu.get('institucion'):
                parrafo(edu['institucion'], tam=TAM_META, cursiva=True,
                        color=COLOR_SECUNDARIO, espacio_despues=2)

    # Las habilidades van en una línea separada por comas: es el formato que
    # mejor indexan los ATS al extraer palabras clave.
    for clave, titulo in (('habilidades_tecnicas', 'Habilidades técnicas'),
                          ('habilidades_blandas', 'Habilidades blandas')):
        if datos.get(clave):
            titulo_seccion(titulo)
            parrafo(' • '.join(datos[clave]), espacio_despues=2)

    if datos.get('certificaciones'):
        titulo_seccion('Certificaciones')
        for cert in datos['certificaciones']:
            vineta(cert)

    if datos.get('idiomas'):
        titulo_seccion('Idiomas')
        parrafo(' • '.join(datos['idiomas']), espacio_despues=2)

    if datos.get('referencias'):
        titulo_seccion('Referencias')
        for ref in datos['referencias']:
            partes = [ref.get('nombre', ''), ref.get('cargo', ''), ref.get('telefono', '')]
            texto = ' — '.join([x for x in partes if x])
            if texto:
                parrafo(texto, espacio_despues=1)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _rango_fechas(exp: Dict[str, Any]) -> str:
    """Devuelve 'inicio - fin' con lo que haya disponible."""
    inicio, fin = exp.get('inicio', ''), exp.get('fin', '')
    if inicio and fin:
        return f'{inicio} - {fin}'
    return inicio or fin or ''


# ==================== PREVISUALIZACIÓN HTML ====================

def _escapar(texto: Any) -> str:
    from html import escape
    return escape(str(texto or ''))


def construir_preview_html(datos: Dict[str, Any]) -> str:
    """Genera la vista previa que se muestra en el navegador.

    Replica el formato del DOCX (misma fuente, mismas proporciones) para que
    lo que el usuario ve sea lo que descarga.
    """
    gris = f'#{COLOR_SECUNDARIO[0]:02x}{COLOR_SECUNDARIO[1]:02x}{COLOR_SECUNDARIO[2]:02x}'
    partes = [
        f'<div class="cv-preview" style="font-family: {FUENTE}, Arial, sans-serif; '
        f'color: #1a1a1a; line-height: {INTERLINEADO}; background: #fff; padding: 20px 18px;">'
    ]

    if datos.get('nombre'):
        partes.append(
            f'<div style="text-align:center; font-size:{TAM_NOMBRE}px; font-weight:700; '
            f'letter-spacing:1.2px;">{_escapar(datos["nombre"]).upper()}</div>'
        )
    if datos.get('titular'):
        partes.append(
            f'<div style="text-align:center; font-size:{TAM_TITULAR}px; color:{gris}; '
            f'margin-top:2px;">{_escapar(datos["titular"])}</div>'
        )

    contacto = [datos.get('contacto', {}).get(c) for c in ('email', 'telefono', 'ciudad', 'linkedin')]
    contacto = [_escapar(c) for c in contacto if c]
    if contacto:
        partes.append(
            f'<div style="text-align:center; font-size:{TAM_CONTACTO}px; color:{gris}; '
            f'padding-bottom:6px; border-bottom:2px solid #444; margin-bottom:2px;">'
            f'{"  •  ".join(contacto)}</div>'
        )

    def titulo(texto):
        return (
            f'<div style="font-size:{TAM_SECCION}px; font-weight:700; text-transform:uppercase; '
            f'letter-spacing:0.8px; border-bottom:1px solid #9a9a9a; margin:14px 0 6px; '
            f'padding-bottom:3px;">{_escapar(texto)}</div>'
        )

    def encabezado_con_fecha(izquierda, derecha):
        """Título a la izquierda y fecha a la derecha, como el tabulador del DOCX."""
        return (
            f'<div style="display:flex; justify-content:space-between; align-items:baseline; '
            f'gap:10px; margin-top:8px;">'
            f'<span style="font-weight:700; font-size:{TAM_CUERPO}px;">{_escapar(izquierda)}</span>'
            f'<span style="font-size:{TAM_META}px; color:{gris}; white-space:nowrap;">'
            f'{_escapar(derecha)}</span></div>'
        )

    if datos.get('resumen'):
        partes.append(titulo('Perfil profesional'))
        partes.append(
            f'<p style="margin:0 0 4px; font-size:{TAM_CUERPO}px; text-align:justify;">'
            f'{_escapar(datos["resumen"])}</p>'
        )

    if datos.get('experiencia'):
        partes.append(titulo('Experiencia profesional'))
        for exp in datos['experiencia']:
            partes.append(encabezado_con_fecha(
                exp.get('cargo') or exp.get('empresa', ''), _rango_fechas(exp)))

            subtitulo = [x for x in (exp.get('empresa') if exp.get('cargo') else '',
                                     exp.get('ubicacion')) if x]
            if subtitulo:
                partes.append(
                    f'<div style="font-size:{TAM_META}px; color:{gris}; font-style:italic;">'
                    f'{_escapar(" — ".join(subtitulo))}</div>'
                )

            if exp.get('logros'):
                items = ''.join(f'<li style="margin-bottom:2px;">{_escapar(l)}</li>'
                                for l in exp['logros'])
                partes.append(
                    f'<ul style="margin:4px 0 0 16px; padding:0; font-size:{TAM_CUERPO}px;">{items}</ul>'
                )

    if datos.get('educacion'):
        partes.append(titulo('Educación'))
        for edu in datos['educacion']:
            partes.append(encabezado_con_fecha(edu.get('titulo', ''), edu.get('anio', '')))
            if edu.get('institucion'):
                partes.append(
                    f'<div style="font-size:{TAM_META}px; color:{gris}; font-style:italic;">'
                    f'{_escapar(edu["institucion"])}</div>'
                )

    for clave, nombre in (('habilidades_tecnicas', 'Habilidades técnicas'),
                          ('habilidades_blandas', 'Habilidades blandas'),
                          ('idiomas', 'Idiomas')):
        if datos.get(clave):
            partes.append(titulo(nombre))
            partes.append(
                f'<p style="margin:0; font-size:{TAM_CUERPO}px;">'
                f'{_escapar(" • ".join(datos[clave]))}</p>'
            )

    if datos.get('certificaciones'):
        partes.append(titulo('Certificaciones'))
        items = ''.join(f'<li>{_escapar(c)}</li>' for c in datos['certificaciones'])
        partes.append(f'<ul style="margin:4px 0 0 16px; padding:0; font-size:{TAM_CUERPO}px;">{items}</ul>')

    if datos.get('referencias'):
        partes.append(titulo('Referencias'))
        for ref in datos['referencias']:
            texto = ' — '.join([x for x in (ref.get('nombre'), ref.get('cargo'),
                                            ref.get('telefono')) if x])
            if texto:
                partes.append(f'<div style="font-size:{TAM_CUERPO}px;">{_escapar(texto)}</div>')

    partes.append('</div>')
    return ''.join(partes)


def resumen_formato() -> List[str]:
    """Reglas de formato aplicadas, para mostrárselas al usuario."""
    return [
        f'Fuente {FUENTE}: {TAM_CUERPO}pt en el cuerpo, {TAM_NOMBRE}pt en el nombre',
        f'Márgenes de {MARGEN_CM} cm e interlineado {INTERLINEADO}',
        'Sin tablas, columnas, imágenes ni cuadros de texto (los ATS no los leen)',
        'Sin encabezado ni pie de página: los datos de contacto van en el cuerpo',
        'Títulos de sección en mayúscula sostenida, con nombres estándar y línea separadora',
        'Cargo en negrita con las fechas alineadas a la derecha, y la empresa en cursiva debajo',
        'Viñetas nativas de Word, no guiones escritos a mano',
        'Habilidades en una línea para facilitar la extracción de palabras clave',
        'Se omiten los datos de terceros de cada empresa (jefe, recursos humanos, NIT)',
    ]


def _sin_tildes(texto: str) -> str:
    return texto.lower().translate(str.maketrans('áéíóúü', 'aeiouu'))


def _verificar_contra_original(datos: Dict[str, Any], cv_text: str) -> Dict[str, Any]:
    """Corrige el nombre del titular si el modelo eligió mal.

    Un fallo observado: en CV con sección de referencias, el modelo toma como
    titular el nombre de una de las referencias. Como el CV acaba en manos de
    un reclutador, un nombre equivocado lo invalida entero, así que se contrasta
    con lo que detecta el análisis por texto.
    """
    heuristico = estructurar_sin_ia(cv_text)
    nombre_modelo = datos.get('nombre', '').strip()

    # Nombres de las referencias, que nunca deben acabar en la cabecera
    nombres_referencias = {_sin_tildes(r['nombre']) for r in heuristico.get('referencias', [])
                           if r.get('nombre')}

    problema = None
    if not nombre_modelo:
        problema = 'el modelo no devolvió nombre'
    elif _sin_tildes(nombre_modelo) in nombres_referencias:
        problema = f'"{nombre_modelo}" es una de las referencias del CV'
    elif _sin_tildes(nombre_modelo) not in _sin_tildes(cv_text):
        problema = f'"{nombre_modelo}" no aparece en el CV original'

    if problema and heuristico.get('nombre'):
        logger.warning(f'Nombre corregido ({problema}): se usa "{heuristico["nombre"]}"')
        datos['nombre'] = heuristico['nombre']

    # Si el modelo se saltó referencias que sí estaban, se conservan las detectadas
    if heuristico.get('referencias') and len(datos.get('referencias', [])) < len(heuristico['referencias']):
        datos['referencias'] = heuristico['referencias']

    # El modelo a veces devuelve solo el primer empleo y descarta el resto.
    # Perder puestos deja al candidato con menos trayectoria de la que tiene,
    # así que los que falten se añaden tal como los detectó el análisis por
    # texto: sin reescribir, pero presentes.
    empresas_modelo = {_sin_tildes(e.get('empresa', '')) for e in datos.get('experiencia', [])}
    cargos_modelo = {_sin_tildes(e.get('cargo', '')) for e in datos.get('experiencia', [])}

    faltantes = [
        exp for exp in heuristico.get('experiencia', [])
        if _sin_tildes(exp.get('empresa', '')) not in empresas_modelo
        and _sin_tildes(exp.get('cargo', '')) not in cargos_modelo
    ]
    if faltantes:
        nombres = ', '.join(e.get('empresa') or e.get('cargo') for e in faltantes)
        logger.warning(f'El modelo omitió {len(faltantes)} puesto(s); se recuperan: {nombres}')
        datos.setdefault('experiencia', []).extend(faltantes)

    return datos


def generar_cv_optimizado(bot, cv_text: str, cargo_objetivo: str = '') -> Dict[str, Any]:
    """Punto de entrada: reestructura el CV y devuelve datos + preview.

    Usa Ollama si está disponible; si no, cae al troceo heurístico para que la
    función siga entregando un documento bien formateado.
    """
    if not cv_text or not cv_text.strip():
        raise ValueError('No hay texto de CV para procesar')

    modo = 'ia'
    if bot is not None and bot.check_ollama_status():
        try:
            datos = estructurar_con_ollama(bot, cv_text, cargo_objetivo)
            datos = _verificar_contra_original(datos, cv_text)
        except Exception as e:
            logger.warning(f'Ollama falló al reestructurar el CV, usando fallback: {e}')
            datos = estructurar_sin_ia(cv_text)
            modo = 'basico'
    else:
        datos = estructurar_sin_ia(cv_text)
        modo = 'basico'

    # Si el modelo devolvió algo vacío, el fallback al menos conserva el contenido
    if not datos.get('experiencia') and not datos.get('resumen') and modo == 'ia':
        logger.warning('Ollama devolvió un CV vacío, usando fallback heurístico')
        datos = estructurar_sin_ia(cv_text)
        modo = 'basico'

    return {
        'datos': datos,
        'preview_html': construir_preview_html(datos),
        'reglas': resumen_formato(),
        'modo': modo,
    }
